"""
Document Processing Service - Extract text from various file formats.
Supports: PDF, DOCX, CSV, TXT, Images (OCR).
"""

import os
from typing import List, Dict, Any, Optional
from pathlib import Path
import logging

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from app.config import settings

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Process and extract text from various document types."""

    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.CHUNK_SIZE,
            chunk_overlap=settings.CHUNK_OVERLAP,
            length_function=len,
            separators=["\n\n", "\n", " ", ""],
        )

    def process_file(
        self,
        file_path: str,
        file_type: str,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> List[Document]:
        """
        Process a file and return document chunks.

        Args:
            file_path: Path to the file
            file_type: Type of file (pdf, docx, csv, txt, image)
            metadata: Additional metadata to attach to documents

        Returns:
            List of Document objects with chunks
        """
        try:
            # Extract text based on file type
            if file_type == "pdf":
                text = self._extract_pdf(file_path)
            elif file_type in ["docx", "doc"]:
                text = self._extract_docx(file_path)
            elif file_type == "csv":
                text = self._extract_csv(file_path)
            elif file_type == "txt":
                text = self._extract_txt(file_path)
            elif file_type in ["png", "jpg", "jpeg"]:
                text = self._extract_image(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")

            # Create base metadata
            base_metadata = {
                "source": file_path,
                "file_type": file_type,
                "file_name": Path(file_path).name,
            }
            if metadata:
                base_metadata.update(metadata)

            # Create document and split into chunks
            doc = Document(page_content=text, metadata=base_metadata)
            chunks = self.text_splitter.split_documents([doc])

            # Add chunk index to metadata
            for idx, chunk in enumerate(chunks):
                chunk.metadata["chunk_index"] = idx
                chunk.metadata["total_chunks"] = len(chunks)

            logger.info(f"Processed {file_path}: {len(chunks)} chunks created")
            return chunks

        except Exception as e:
            logger.error(f"Error processing file {file_path}: {str(e)}")
            raise

    def _extract_pdf(self, file_path: str) -> str:
        """Extract text from PDF file."""
        try:
            from PyPDF2 import PdfReader

            reader = PdfReader(file_path)
            text_parts = []

            for page_num, page in enumerate(reader.pages, start=1):
                text = page.extract_text()
                if text.strip():
                    text_parts.append(f"[Page {page_num}]\n{text}")

            return "\n\n".join(text_parts)

        except Exception as e:
            logger.error(f"Error extracting PDF: {str(e)}")
            raise ValueError(f"Failed to extract PDF: {str(e)}")

    def _extract_docx(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(file_path)
            text_parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)

            return "\n\n".join(text_parts)

        except Exception as e:
            logger.error(f"Error extracting DOCX: {str(e)}")
            raise ValueError(f"Failed to extract DOCX: {str(e)}")

    def _extract_csv(self, file_path: str) -> str:
        """Extract text from CSV file."""
        try:
            import pandas as pd

            # Try to read CSV with different encodings
            encodings = ["utf-8", "latin-1", "iso-8859-1", "cp1252"]
            df = None

            for encoding in encodings:
                try:
                    df = pd.read_csv(file_path, encoding=encoding)
                    break
                except UnicodeDecodeError:
                    continue

            if df is None:
                raise ValueError("Could not read CSV with any encoding")

            # Convert to text representation
            text_parts = [f"CSV File: {Path(file_path).name}"]
            text_parts.append(f"Columns: {', '.join(df.columns.tolist())}")
            text_parts.append(f"Rows: {len(df)}")
            text_parts.append("\n--- Data ---\n")

            # Add summary statistics for numeric columns
            numeric_cols = df.select_dtypes(include=["number"]).columns
            if len(numeric_cols) > 0:
                text_parts.append("Summary Statistics:")
                text_parts.append(df[numeric_cols].describe().to_string())
                text_parts.append("")

            # Add first few rows
            text_parts.append("Sample Data (first 20 rows):")
            text_parts.append(df.head(20).to_string())

            return "\n".join(text_parts)

        except Exception as e:
            logger.error(f"Error extracting CSV: {str(e)}")
            raise ValueError(f"Failed to extract CSV: {str(e)}")

    def _extract_txt(self, file_path: str) -> str:
        """Extract text from TXT file."""
        try:
            # Try different encodings
            encodings = ["utf-8", "latin-1", "iso-8859-1", "cp1252"]

            for encoding in encodings:
                try:
                    with open(file_path, "r", encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue

            raise ValueError("Could not read file with any encoding")

        except Exception as e:
            logger.error(f"Error extracting TXT: {str(e)}")
            raise ValueError(f"Failed to extract TXT: {str(e)}")

    def _extract_image(self, file_path: str) -> str:
        """Extract text from image using OCR."""
        try:
            import pytesseract
            from PIL import Image

            # Open image
            image = Image.open(file_path)

            # Perform OCR
            text = pytesseract.image_to_string(image)

            if not text.strip():
                return f"[Image: {Path(file_path).name}] - No text detected"

            return f"[Image: {Path(file_path).name}]\n\n{text}"

        except ImportError:
            logger.warning("pytesseract not installed, returning image metadata only")
            return f"[Image: {Path(file_path).name}] - OCR not available"

        except Exception as e:
            logger.error(f"Error extracting image text: {str(e)}")
            return f"[Image: {Path(file_path).name}] - OCR failed: {str(e)}"


# Singleton instance
document_processor = DocumentProcessor()
